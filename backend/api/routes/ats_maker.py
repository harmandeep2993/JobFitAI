# api/routes/ats_maker.py
"""
/api/ats/* endpoints - ATS check, optimise, and DOCX export.
"""

import io

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException
from fastapi.concurrency import run_in_threadpool
from fastapi.responses import JSONResponse, StreamingResponse

from api.routes.auth import get_current_user, get_current_user_llm_limited
from core import uploads
from core.config import JD_MAX_CHARS
from core.logger import get_logger
from repositories import resume_repo as resume_store
from schemas.ats import AtsCheckRequest, AtsDocxRequest, AtsOptimiseRequest
from services.ats import ats_check, generate_ats_resume
from services.extractors.jd_extractor import extract_jd
from services.parsers import extract_all_text

logger = get_logger(__name__)

router = APIRouter()

# Body text size for generated DOCX files - standard resume body size.
_DOCX_BODY_PT = 11

# The candidate's name leads the document and should read as the title.
_DOCX_NAME_PT = 16

# Gap between columns in twips (1/20 pt). 432 = 0.3 inch.
_COLUMN_GAP_TWIPS = 432


def _resolve_resume_text(user_id: str, resume_id: str, tmp: str) -> str:
    """Return extracted text for a stored resume or a temp upload token.

    Raises HTTPException when neither source resolves to readable text.
    """
    path = None
    if resume_id:
        record = resume_store.get(user_id, resume_id)
        if not record:
            raise HTTPException(status_code=404, detail="resume_not_found")
        path = record["file_path"]
    elif tmp:
        path = uploads.resolve(user_id, tmp)
        if not path:
            raise HTTPException(status_code=404, detail="resume_not_found")
    else:
        raise HTTPException(status_code=400, detail="resume_required")

    text = extract_all_text(path)
    if not text:
        raise HTTPException(status_code=422, detail="could_not_read_resume")
    return text


@router.post("/check")
async def api_ats_check(
    body: AtsCheckRequest,
    current_user: dict = Depends(get_current_user_llm_limited),
) -> JSONResponse:
    """
    Lightweight ATS scan - no LLM, no keyword injection.

    Accepts {resume_id | tmp, jd?}. Returns section heading flags, formatting
    flags, keyword coverage (if JD provided), and a composite ATS score.
    """
    resume_text = await run_in_threadpool(
        _resolve_resume_text,
        current_user["id"],
        (body.resume_id or "").strip(),
        (body.tmp or "").strip(),
    )

    jd_text = (body.jd or "").strip()[:JD_MAX_CHARS]
    required_skills = None
    if jd_text and len(jd_text) >= 50:
        jd_json = await run_in_threadpool(extract_jd, jd_text)
        if jd_json:
            required_skills = jd_json.get("required_skills") or []

    result = await run_in_threadpool(ats_check, resume_text, required_skills)
    return JSONResponse({"ok": True, **result})


@router.post("/optimise")
async def api_ats_optimise(
    body: AtsOptimiseRequest,
    current_user: dict = Depends(get_current_user_llm_limited),
) -> JSONResponse:
    """
    Full ATS optimization pipeline with LLM.

    Accepts {resume_id | tmp, jd}. Returns a complete ATS-optimised resume
    with coverage before/after, section flags, and formatting warnings.
    """
    jd_text = (body.jd or "").strip()[:JD_MAX_CHARS]
    if len(jd_text) < 50:
        raise HTTPException(status_code=400, detail="jd_required")

    resume_text = await run_in_threadpool(
        _resolve_resume_text,
        current_user["id"],
        (body.resume_id or "").strip(),
        (body.tmp or "").strip(),
    )

    result = await run_in_threadpool(generate_ats_resume, resume_text, jd_text)
    if not result:
        raise HTTPException(status_code=503, detail="llm_unavailable")

    return JSONResponse({"ok": True, **result})


def _set_column_count(section, count: int) -> None:
    """Set a section's column count via the raw sectPr XML.

    python-docx has no API for this, but the underlying w:cols element does. This
    is how you get a two-column resume that ATS parsers can still read: Word
    columns are a RENDERING instruction, so the paragraphs stay in linear order
    in the XML and text extraction walks them top to bottom. A table, by
    contrast, builds a grid that parsers read cell by cell - which is exactly how
    two-column resumes end up with the candidate's job titles interleaved with
    their skills, or dropped entirely.
    """
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(_COLUMN_GAP_TWIPS))


def _add_section_heading(doc, text: str) -> None:
    """Standard ATS heading - the exact words parsers look for."""
    doc.add_heading(text, level=1)


def _add_experience(doc, experience: list) -> None:
    for job in experience:
        header = " | ".join(
            p for p in [job.get("title"), job.get("company"), job.get("dates")] if p
        )
        p = doc.add_paragraph()
        p.add_run(header).bold = True
        for b in job.get("bullets") or []:
            doc.add_paragraph(b, style="List Bullet")


def _add_education(doc, education: list) -> None:
    for edu in education:
        line = " - ".join(p for p in [edu.get("degree"), edu.get("institution")] if p)
        year = edu.get("year")
        doc.add_paragraph(f"{line} ({year})" if year else line)


def _build_docx(parsed: dict, layout: str = "single") -> bytes:
    """Render the optimised-resume JSON into a sendable, ATS-friendly DOCX.

    Single column, standard headings, no tables or text boxes - the layouts that
    make ATS parsers scramble or skip content. The contact block is plain text at
    the top (never a header/footer or image), because that is the single most
    common reason a resume parses with no name or email attached.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(_DOCX_BODY_PT)

    # === Contact header ===
    # A resume without one is not sendable, and an ATS that cannot find an email
    # has nothing to attach the application to.
    contact = parsed.get("contact") or {}
    name = (contact.get("name") or "").strip()
    if name:
        heading = doc.add_paragraph()
        run = heading.add_run(name)
        run.bold = True
        run.font.size = Pt(_DOCX_NAME_PT)

    contact_bits = [
        (contact.get("email") or "").strip(),
        (contact.get("phone") or "").strip(),
        (contact.get("location") or "").strip(),
    ]
    contact_bits += [
        link.strip()
        for link in (contact.get("links") or [])
        if isinstance(link, str) and link.strip()
    ]
    contact_bits = [b for b in contact_bits if b]
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))

    summary = (parsed.get("summary") or "").strip()
    if summary:
        _add_section_heading(doc, "Summary")
        doc.add_paragraph(summary)

    experience = parsed.get("experience") or []
    skills = parsed.get("skills") or []
    education = parsed.get("education") or []

    if layout == "two_column":
        # The header and summary stay full width (a recruiter's eye and an ATS
        # both want the name and contact block unambiguous). Everything below
        # runs in a two-column band.
        #
        # Reading order is what makes this safe: the paragraphs are emitted
        # sidebar-first, then a COLUMN break, then the main content. Word renders
        # that as two columns; a text extractor walks the same paragraphs in the
        # same order and gets Skills, Education, then Work Experience - each
        # section intact and under its own standard heading.
        body = doc.add_section(WD_SECTION.CONTINUOUS)
        _set_column_count(body, 2)

        if skills:
            _add_section_heading(doc, "Skills")
            for skill in skills:
                doc.add_paragraph(skill, style="List Bullet")

        if education:
            _add_section_heading(doc, "Education")
            _add_education(doc, education)

        # Push the main content into the second column.
        doc.add_paragraph().add_run().add_break(WD_BREAK.COLUMN)

        if experience:
            _add_section_heading(doc, "Work Experience")
            _add_experience(doc, experience)
    else:
        if experience:
            _add_section_heading(doc, "Work Experience")
            _add_experience(doc, experience)

        if skills:
            _add_section_heading(doc, "Skills")
            doc.add_paragraph(", ".join(skills))

        if education:
            _add_section_heading(doc, "Education")
            _add_education(doc, education)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/docx")
async def api_ats_docx(
    body: AtsDocxRequest,
    current_user: dict = Depends(get_current_user),
) -> StreamingResponse:
    """Download the optimised resume (JSON from /optimise) as a DOCX file."""
    if not body.resume:
        raise HTTPException(status_code=400, detail="resume_required")

    layout = (body.layout or "single").strip().lower()
    if layout not in ("single", "two_column"):
        raise HTTPException(status_code=400, detail="invalid_layout")

    data = await run_in_threadpool(_build_docx, body.resume, layout)
    return StreamingResponse(
        iter([data]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=ats_optimised_resume.docx"
        },
    )
